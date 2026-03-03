import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


def setup_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(0)
    return doc


def add_heading(doc: Document, text: str, level: int) -> None:
    level = max(0, min(level, 4))
    p = doc.add_heading(text.strip(), level=level)
    p.paragraph_format.line_spacing = 1.5


def add_code_block(doc: Document, lines):
    if not lines:
        return
    p = doc.add_paragraph("\n".join(lines))
    p.style = doc.styles["Normal"]
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    run = p.runs[0]
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
    run.font.size = Pt(10.5)


def parse_table(lines):
    rows = []
    for line in lines:
        line = line.strip()
        if not line.startswith("|"):
            continue
        cells = [c.strip() for c in line.strip("|").split("|")]
        rows.append(cells)
    if len(rows) < 2:
        return None
    # drop separator row (---)
    header = rows[0]
    body = [r for r in rows[1:] if not all(re.fullmatch(r":?-{3,}:?", c or "") for c in r)]
    return header, body


def add_table(doc: Document, header, body):
    cols = len(header)
    table = doc.add_table(rows=1, cols=cols)
    table.style = "Table Grid"
    for i, h in enumerate(header):
        table.rows[0].cells[i].text = h
    for row in body:
        cells = table.add_row().cells
        for i in range(cols):
            cells[i].text = row[i] if i < len(row) else ""


def convert(md_path: Path, docx_path: Path) -> None:
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()
    doc = setup_document()

    in_code = False
    code_lines = []
    table_buffer = []

    def flush_table():
        nonlocal table_buffer
        if table_buffer:
            parsed = parse_table(table_buffer)
            if parsed:
                add_table(doc, parsed[0], parsed[1])
            else:
                for raw in table_buffer:
                    if raw.strip():
                        doc.add_paragraph(raw)
            table_buffer = []

    for raw in lines:
        line = raw.rstrip("\n")

        if in_code:
            if line.strip().startswith("```"):
                add_code_block(doc, code_lines)
                in_code = False
                code_lines = []
            else:
                code_lines.append(line)
            continue

        if line.strip().startswith("```"):
            flush_table()
            in_code = True
            code_lines = []
            continue

        if line.strip().startswith("|"):
            table_buffer.append(line)
            continue
        else:
            flush_table()

        if not line.strip():
            doc.add_paragraph("")
            continue

        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            level = len(m.group(1)) - 1
            add_heading(doc, m.group(2), level)
            continue

        if re.match(r"^\s*[-*]\s+", line):
            text_line = re.sub(r"^\s*[-*]\s+", "", line)
            doc.add_paragraph(text_line, style="List Bullet")
            continue

        if re.match(r"^\s*\d+\.\s+", line):
            text_line = re.sub(r"^\s*\d+\.\s+", "", line)
            doc.add_paragraph(text_line, style="List Number")
            continue

        p = doc.add_paragraph(line)
        p.paragraph_format.first_line_indent = Cm(0.74)

    flush_table()

    # title alignment for first heading if exists
    for para in doc.paragraphs:
        if para.style.name.startswith("Heading") and para.text.strip():
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            break

    doc.save(str(docx_path))


def main():
    if len(sys.argv) < 3:
        print("Usage: python scripts/md_to_docx.py <input.md> <output.docx>")
        sys.exit(1)
    md_path = Path(sys.argv[1])
    docx_path = Path(sys.argv[2])
    if not md_path.exists():
        print(f"Input not found: {md_path}")
        sys.exit(2)
    convert(md_path, docx_path)
    print(f"Generated: {docx_path}")


if __name__ == "__main__":
    main()
